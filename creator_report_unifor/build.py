# import streamlit as st
from docx import Document
import pandas as pd
import streamlit as st

st.header('Conversor Semgrep Relatorio')

my_file = st.file_uploader(label="XLSX from SEMGREP", type="xlsx")

if st.button(label="Gerar Relatorio"):
    if my_file:

        with open(my_file.name, 'wb') as file:
            file.write(my_file.getvalue())

        data_semgrep = pd.read_excel(my_file.name).to_dict(orient='records')

        doc = Document()
        doc.add_picture(image_path_or_stream="unifor-logo.jpg")
        doc.add_heading(text="Semgrep Relatório de Vulnerabilidades", level=1)
        doc.add_page_break()
        for vuln in data_semgrep:
            doc.add_heading(vuln.get('Check_id'), level=2)
            pr = doc.add_paragraph()
            pr.add_run(text="Arquivo: ").bold = True
            pr.add_run(text=f"{vuln.get('Path')}\n")

            pr.add_run(text="Severidade: ").bold = True
            pr.add_run(text=f"{vuln.get('Severity')}\n")

            pr.add_run(text="Linha do Código: ").bold = True
            pr.add_run(text=f"{vuln.get('Start_line')}\n")

            pr.add_run(text="Descrição da Vulnerabilidade: ").bold = True
            pr.add_run(text=f"{vuln.get('Message')}\n")

        doc.save(my_file.name.replace('.xlsx', '.docx'))
        st.balloons()
        with open(my_file.name.replace('.xlsx', '.docx'), "rb") as file:
            btn = st.download_button(
                label="Download Document",
                data=file,
                file_name=my_file.name.replace('.xlsx', '.docx'),
            )